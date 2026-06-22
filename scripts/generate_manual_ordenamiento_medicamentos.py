"""Genera Manual_Consumo_Ordenamiento_Medicamentos.docx en la raíz del proyecto."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row_data in rows:
        row = table.add_row().cells
        for i, cell in enumerate(row_data):
            row[i].text = cell
    doc.add_paragraph("")


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    out = root / "Manual_Consumo_Ordenamiento_Medicamentos.docx"

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading(
        "Manual de consumo — Ordenamiento y autorización de medicamentos",
        0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "API ORQUESTADORDB (EPS Familiar de Colombia). "
        "Integración con Messiah (PostgreSQL) para solicitud y autorización automática "
        "de medicamentos por orden médica ambulatoria."
    )
    doc.add_paragraph(
        "Versión del documento: generada desde el código del repositorio. "
        "Consulte también GET /docs y GET /openapi.json en el servidor desplegado."
    )

    doc.add_heading("1. Resumen del servicio", level=1)
    doc.add_paragraph(
        "El ordenamiento de medicamentos permite registrar una solicitud de medicamentos prescritos "
        "por un médico, evaluar reglas de negocio (afiliado activo, contrato/tarifario IPS, "
        "direccionamiento, catálogo CUM, topes) y devolver autorización total, parcial o ninguna, "
        "con PIN de activación cuando aplica."
    )
    doc.add_paragraph(
        "Flujo IPS en dos pasos (Messiah authorization_request_ips + SsAutorizacionActivacion):\n"
        "• Paso 1 — POST /afiliados/autorizacion-orden-medica-ips: solicitud, evaluación y autorización "
        "emitida (estado EMITIDA) con PIN; autorizacion_activa=false; PDF código activación (base64).\n"
        "• Paso 2 — POST /afiliados/autorizacion-orden-medica-ips/activar: activa (ACTIVADA), programa "
        "medicamentos, confirma prestación (COMPLETADA) por defecto; PDF autorización prestada (base64, PRESTADO=1)."
    )

    doc.add_heading("2. Requisitos previos", level=1)
    doc.add_paragraph(
        "2.1 Autenticación JWT\n"
        "Obtener token con POST /auth/login (application/json). "
        "En todas las peticiones de medicamentos enviar: Authorization: Bearer <access_token>."
    )
    doc.add_paragraph(
        "2.2 Permiso autoriza_med\n"
        "El usuario API debe tener orq.usuarios.autoriza_med = 1 (true) en SQL Server (OrquestacionDB). "
        "Sin este flag la API responde 403."
    )
    doc.add_paragraph(
        "Ejemplo SQL:\n"
        "UPDATE orq.usuarios SET autoriza_med = 1 WHERE username = 'su_usuario_api';"
    )
    doc.add_paragraph(
        "2.3 Afiliado\n"
        "El afiliado debe existir en Messiah y estar en estado ACTIVO (código 1) para autorización IPS."
    )
    doc.add_paragraph(
        "2.4 Datos parametrizados en Messiah\n"
        "IPS en CT_IPS, sede en CT_IPS_SEDE, médico en TB_MEDICO_SOLICITANTE, diagnóstico en TB_CIE10, "
        "medicamento en TB_MEDICAMENTO, contrato/tarifario vigente para el municipio del afiliado, "
        "y direccionamiento del CUM en la IPS de autorización cuando la IPS lo exige."
    )

    doc.add_heading("3. URL base y documentación", level=1)
    doc.add_paragraph(
        "URL base según despliegue (ejemplo desarrollo: http://localhost:8000).\n"
        "Swagger: GET /docs | OpenAPI: GET /openapi.json | ReDoc: GET /redoc"
    )
    doc.add_paragraph(
        "Colección Postman: postman/EPS_Familiar_de_Colombia_API.postman_collection.json\n"
        "Entorno: postman/EPS_Familiar_de_Colombia_API.postman_environment.json"
    )

    doc.add_heading("4. Flujo en dos pasos", level=1)
    _add_table(
        doc,
        ["Paso", "Ruta", "Resultado"],
        [
            ["1 — Solicitud", "POST /afiliados/autorizacion-orden-medica-ips", "ss_solicitud + autorización EMITIDA, PIN, pendiente_activacion=true"],
            ["2 — Activación y confirmación", "POST /afiliados/autorizacion-orden-medica-ips/activar", "ACTIVADA (2) + COMPLETADA (3), PDF prestado, soporte opcional en autorizacion_confirmada/"],
        ],
    )

    doc.add_heading("5. POST /afiliados/autorizacion-orden-medica-ips (paso 1)", level=1)
    doc.add_paragraph("Ruta: {baseUrl}/afiliados/autorizacion-orden-medica-ips")
    doc.add_paragraph(
        "Integración para IPS externas. El tarifario y la autorización se validan contra "
        "nit_ips_direccionamiento en el municipio del afiliado."
    )

    doc.add_heading("5.1 Campos del formulario", level=2)
    _add_table(
        doc,
        ["Campo", "Obligatorio", "Tipo API", "Descripción"],
        [
            ["tipo_identificacion", "Sí", "string", "Tipo documento afiliado"],
            ["numero_identificacion", "Sí", "string", "Número documento"],
            ["observacion", "Sí", "string", "Texto solicitud"],
            ["email", "No", "string", "Si vacío: sede → afiliado → IPS"],
            ["telefono", "Sí", "string", "Obligatorio. Formato 9999999999"],
            ["celular", "Sí", "string", "Obligatorio. Formato 9999999999"],
            ["prioridad_atencion", "No", "integer", "Default 1"],
            ["justificacion_clinica", "No", "string", ""],
            ["diagnostico_principal", "Sí", "string", "CIE-10 obligatorio"],
            ["diagnostico_relacionado_1", "No", "string", "CIE-10 opcional"],
            ["diagnostico_relacionado_2", "No", "string", "CIE-10 opcional"],
            ["nit_ips_prestador", "Sí", "string", "IPS solicitante"],
            ["nit_ips_direccionamiento", "Sí", "string", "IPS donde se autoriza / tarifario"],
            ["consecutivo_sede_ips", "No", "integer", "Primera sede habilitada si vacío"],
            ["telefono_institucional_extension", "No", "string", ""],
            ["prestador_solicitante", "No", "string", "Opcional si envía NIT"],
            ["municipio_prestador", "No", "string", "Autocompleta desde sede/IPS"],
            ["fecha_solicitud_proceso", "Sí", "date", "YYYY-MM-DD, ≤ hoy"],
            ["fecha_solicitud_medico", "Sí", "date", "YYYY-MM-DD, ≤ hoy"],
            ["registro_profesional", "Sí", "string", "tb_medico_solicitante"],
            ["medicamentos_json", "Sí", "string (JSON)", "Solo cum, cantidad, dias, observacion"],
            ["soporte_orden_medica", "Sí", "file (PDF)", "Máx. ORDEN_MEDICA_SOPORTE_MAX_MB (default 5)"],
        ],
    )

    doc.add_heading("5.2 Estructura medicamentos_json", level=2)
    doc.add_paragraph(
        'Ejemplo:\n'
        '[{"cum":"13874-1","cantidad":1,"dias":30,"observacion":"Uso segun orden medica"}]\n\n'
        "No enviar campos extra (posología se resuelve del catálogo tb_medicamento). "
        "No repetir el mismo CUM en la misma solicitud."
    )

    doc.add_heading("5.3 Valores fijados por el servidor", level=2)
    doc.add_paragraph(
        "• origen_solicitud: ORDEN_MEDICA\n"
        "• ubicacion_paciente: AMBULATORIO\n"
        "• Origen de atención en BD: Enfermedad General (ss_solicitud_atencion)\n"
        "• Modalidad: Ambulatorios (tb_modalidad_servicio_salud por descripción/código)\n"
        "• Contacto institucional: desde CT_IPS_SEDE / afiliado / IPS"
    )

    doc.add_heading("6. POST /afiliados/autorizacion-orden-medica-ips/activar (paso 2)", level=1)
    doc.add_paragraph("Ruta: {baseUrl}/afiliados/autorizacion-orden-medica-ips/activar")
    doc.add_paragraph(
        "Activa y confirma la prestación en una sola llamada (Messiah SsAutorizacionActivacion). "
        "Programa medicamentos (fecha_programacion), activa (estado ACTIVADA=2), confirma prestación "
        "(COMPLETADA=3) y devuelve PDF autorización prestada en base64 (PRESTADO=1). "
        "Idempotente si ya estaba activada o confirmada. Compatible con solicitudes creadas en Messiah o en este API."
    )
    _add_table(
        doc,
        ["Campo", "Obligatorio", "Tipo API", "Descripción"],
        [
            ["tipo_identificacion", "Sí", "string", "Mismo afiliado del paso 1"],
            ["numero_identificacion", "Sí", "string", "Documento del afiliado"],
            ["pin_activacion", "Sí", "string", "PIN devuelto en paso 1 (ss_autorizacion.pin)"],
            ["nit_ips_direccionamiento", "Sí", "string", "Mismo NIT del paso 1"],
            ["fecha_programacion", "No", "date", "Programación medicamentos. Default: hoy (YYYY-MM-DD)"],
            ["fecha_prestacion", "No", "date", "Fecha prestación servicio. Default: hoy (YYYY-MM-DD)"],
            ["confirmar_prestacion", "No", "boolean", "Default true. Si false, solo activa sin confirmar"],
            ["soporte_confirmacion", "No", "file (PDF)", "Opcional. Repositorio Messiah autorizacion_confirmada/"],
        ],
    )

    doc.add_heading("6.1 Respuesta paso 2 (campos principales)", level=2)
    _add_table(
        doc,
        ["Campo", "Descripción"],
        [
            ["prestacion_confirmada", "true cuando se confirmó la prestación"],
            ["estado_trazabilidad", "COMPLETADA (paso 2 con confirmar_prestacion=true) o ACTIVADA"],
            ["fecha_programacion / fecha_prestacion", "Fechas usadas en la operación"],
            ["fecha_real_prestacion_servicio", "Fecha persistida en ss_autorizacion"],
            ["soporte_confirmacion_registrado_messiah", "true si el PDF opcional se guardó en Messiah"],
            ["consecutivo_saldo", "PK sc_saldo_encabezado (nota NC-AT) si preferencias 288/390 contabilizan en activación/confirmación"],
            ["ya_activada / ya_confirmada", "Idempotencia: true si ya existía el estado"],
            ["autorizacion_pdf_base64", "PDF autorización prestada (PRESTADO=1) cuando Jasper está disponible"],
        ],
    )

    doc.add_heading("7. Reglas de autorización automática", level=1)
    doc.add_paragraph(
        "Por cada medicamento el motor evalúa, entre otras:\n"
        "• Existencia y estado activo en tb_medicamento (sw_activo).\n"
        "• Marcación automática en tarifario del contrato (tb_tarifario_propio_detalle.sw_automatico=1).\n"
        "• Al menos un contrato IPS activo con tarifario de medicamentos y cobertura del municipio del afiliado.\n"
        "• telefono y celular obligatorios en la solicitud (validación de formato).\n"
        "• IPS habilitada (sw_habilitada) y tipo_autoriza=2 o tipo_medicamento=2 (autorización automática Messiah).\n"
        "• Direccionamiento del medicamento en la IPS cuando aplica.\n"
        "• Validaciones adicionales Messiah (topes, etc.) vía medicamento_validacion_messiah."
    )
    doc.add_paragraph(
        "tipo_resultado en la respuesta:\n"
        "• Si ningún medicamento cumple validaciones → HTTP 400 y no se persiste solicitud.\n"
        "• TOTAL — todos los ítems autorizados.\n"
        "• PARCIAL — al menos uno autorizado y solicitud creada."
    )

    doc.add_heading("8. PDF en Base64 (Messiah / Jasper)", level=1)
    doc.add_paragraph(
        "Si el servidor tiene Java 8 (Temurin) y JasperStarter (MESSIAH_PDF_ENABLED=true), las respuestas incluyen:\n"
        "• pdf_generado — true cuando se generó el archivo.\n"
        "• autorizacion_pdf_nombre — nombre sugerido del PDF.\n"
        "• autorizacion_pdf_base64 — contenido PDF (RFC 4648).\n"
        "• pdf_aviso — mensaje si Jasper no está disponible o el PDF quedó vacío.\n\n"
        "Paso 2 (confirmación): reporte reAutorizacion.jasper con PRESTADO=1.\n\n"
        "Reportes en app/reports/messiah. En Docker: imagen con Temurin 8 + JASPERSTARTER_PATH=/app/tools/jasperstarter/bin/jasperstarter.\n"
        "La primera generación por autorización compila el reporte (~5–8 s); las siguientes reutilizan caché en JASPER_CACHE_DIR (default /tmp/orq_jasper_cache)."
    )
    doc.add_paragraph(
        "Interoperabilidad: la solicitud puede crearse por este API o por Messiah. El paso 2 consolida activación "
        "y confirmación de prestación. Si ya estaba activada o confirmada en Messiah, POST /activar responde 200 "
        "idempotente con el PDF correspondiente cuando Jasper está disponible.\n\n"
        "Soportes Messiah (MESSIAH_DESCARGAS_RUTA o SFTP):\n"
        "• Paso 1: soporte_ips_solicitud_autorizacion/ y soporte_solicitud/\n"
        "• Paso 2 (obligatorio PDF): autorizacion_confirmada/{consecutivo_autorizacion}-{archivo}.pdf"
    )

    doc.add_heading("8.1 Contabilización (sc_saldo_encabezado / sc_saldo_detalle)", level=2)
    doc.add_paragraph(
        "Según preferencias Messiah 288 (activación) y 390 (confirmación), el paso 2 puede generar la nota "
        "contable de autorización (tipo_documento=1, prefijo NC-AT) y vincular ss_autorizacion.consecutivo_saldo.\n\n"
        "Reglas:\n"
        "• documento en sc_saldo_encabezado = consecutivo_autorizacion (PK numérica).\n"
        "• documento_nota = 'NC-AT {secuencia}' (nota contabilidad autorización).\n"
        "• No se reutilizan saldos de facturación aunque compartan el mismo número en documento.\n"
        "• Si consecutivo_saldo apunta a un asiento que no es NC-AT de esa autorización, se desvincula y se crea uno nuevo.\n"
        "• Concurrencia: bloqueo transaccional pg_advisory_xact_lock en contabilización.\n\n"
        "Tablas: administrativo.sc_saldo_encabezado, sc_saldo_detalle, sc_saldo_encabezado_documento."
    )

    doc.add_heading("8.2 Rendimiento y concurrencia", level=2)
    doc.add_paragraph(
        "• Paso 1: una transacción con lock orq_direccionamiento (advisory) por solicitud.\n"
        "• Paso 2: activación + confirmación + contabilización + PDF en una petición; tiempo típico 8–15 s "
        "(primera vez con PDF ~15–20 s por compilación Jasper).\n"
        "• Reintentos del mismo paso 2 son idempotentes y más rápidos (caché Jasper por consecutivo_autorizacion).\n"
        "• Despliegue: docker compose build && docker compose up -d tras actualizar código.\n"
        "• ss_solicitud: sw_terminada_hospitalaria=1 y servicio='Medicamentos' para que Messiah muestre datos en consulta."
    )

    doc.add_heading("9. Respuesta exitosa (200)", level=1)
    doc.add_paragraph(
        "Paso 1: AutorizacionOrdenMedicaIpsResponse. Paso 2: ActivacionOrdenMedicaIpsResponse. "
        "Campos principales del paso 1:"
    )
    _add_table(
        doc,
        ["Campo", "Descripción"],
        [
            ["consecutivo_solicitud", "ID solicitud en Messiah (ss_solicitud)"],
            ["consecutivo_solicitud_ips", "Registro en orq (solicitud usuario IPS)"],
            ["numero_solicitud", "Número visible de solicitud"],
            ["tipo_resultado", "TOTAL | PARCIAL (si hay solicitud); NINGUNA solo en HTTP 400"],
            ["pin_activacion", "PIN para el paso 2 (si hay autorización)"],
            ["pendiente_activacion", "true tras paso 1; false tras activar"],
            ["autorizacion_activa", "false tras paso 1; true tras activar"],
            ["estado_trazabilidad", "EMITIDA (paso 1) / COMPLETADA (paso 2)"],
            ["consecutivo_autorizacion", "ID autorización generada"],
            ["medicamentos_solicitados", "Listado con estado por ítem"],
            ["medicamentos_autorizados", "Ítems aprobados"],
            ["medicamentos_no_autorizados", "Ítems rechazados con motivo"],
            ["prestador_solicitante / prestador_direccionamiento", "Resumen NIT y razón social (IPS)"],
            ["medico_solicitante", "Datos del médico (endpoint IPS)"],
            ["ips_autorizada", "IPS, municipio, dirección autorizados"],
            ["cobro", "Información de cobro asociada"],
        ],
    )

    doc.add_heading("10. Códigos de error habituales", level=1)
    _add_table(
        doc,
        ["HTTP", "Causa típica"],
        [
            ["401", "Token ausente, inválido o expirado"],
            ["403", "autoriza_med = 0 o sin permiso en seg.usuario_endpoints"],
            ["400", "Afiliado no activo, fechas futuras, PDF inválido, reglas de negocio o ningún medicamento autorizado (sin crear solicitud)"],
            ["404", "Afiliado, IPS, médico, CIE-10, medicamento o autorización pendiente no encontrada"],
            ["409", "No se pudo activar o confirmar (estado inconsistente) o fallo contabilización"],
            ["422", "Validación de formulario (fecha no ISO, entero inválido, JSON mal formado)"],
            ["500", "Error interno o SQL (revisar detail en respuesta)"],
        ],
    )

    doc.add_heading("11. Ejemplo de consumo (cURL — IPS)", level=1)
    doc.add_paragraph(
        "1. Login:\n"
        'curl -X POST "{baseUrl}/auth/login" -H "Content-Type: application/json" '
        '-d "{\\"username\\":\\"admin\\",\\"password\\":\\"***\\"}"\n\n'
        "2. Autorización (reemplace TOKEN y rutas de archivo):\n"
        'curl -X POST "{baseUrl}/afiliados/autorizacion-orden-medica-ips" \\\n'
        '  -H "Authorization: Bearer TOKEN" \\\n'
        '  -F "tipo_identificacion=CC" \\\n'
        '  -F "numero_identificacion=3999852" \\\n'
        '  -F "observacion=Solicitud medicamentos orden médica ambulatoria" \\\n'
        '  -F "diagnostico_principal=102593" \\\n'
        '  -F "nit_ips_prestador=901483168" \\\n'
        '  -F "nit_ips_direccionamiento=901483168" \\\n'
        '  -F "fecha_solicitud_proceso=2026-06-01" \\\n'
        '  -F "fecha_solicitud_medico=2026-06-01" \\\n'
        '  -F "registro_profesional=893604" \\\n'
        '  -F \'medicamentos_json=[{"cum":"13874-1","cantidad":1,"dias":30,"observacion":"Uso segun orden medica"}]\' \\\n'
        '  -F "soporte_orden_medica=@orden_medica.pdf;type=application/pdf"\n\n'
        "3. Activación y confirmación (use PIN y NIT del paso 1):\n"
        'curl -X POST "{baseUrl}/afiliados/autorizacion-orden-medica-ips/activar" \\\n'
        '  -H "Authorization: Bearer TOKEN" \\\n'
        '  -F "tipo_identificacion=CC" \\\n'
        '  -F "numero_identificacion=3999852" \\\n'
        '  -F "pin_activacion=ABCD1234" \\\n'
        '  -F "nit_ips_direccionamiento=901483168" \\\n'
        '  -F "confirmar_prestacion=true"\n\n'
        "Con soporte de confirmación opcional:\n"
        '  -F "soporte_confirmacion=@confirmacion.pdf;type=application/pdf"'
    )

    doc.add_heading("12. Variables de entorno del servidor", level=1)
    doc.add_paragraph(
        "• ORDEN_MEDICA_SOPORTE_MAX_MB — tamaño máximo PDF soporte orden médica (default 5).\n"
        "• MESSIAH_PDF_ENABLED — generar PDF Jasper en respuestas (default true).\n"
        "• JASPERSTARTER_PATH — ruta al ejecutable jasperstarter (opcional).\n"
        "• MESSIAH_JASPER_DIR — carpeta de reportes .jrxml/.jasper.\n"
        "• POSTGRES_URL — conexión Messiah.\n"
        "• SQLSERVER_URL — OrquestacionDB (usuarios, logs).\n"
        "• JWT_SECRET, JWT_EXPIRE_MINUTES — autenticación.\n"
        "• MESSIAH_DESCARGAS_RUTA — montaje local al repositorio de descargas Messiah.\n"
        "• MESSIAH_SOPORTE_TRANSPORT — auto | local | sftp para soportes paso 1 y 2."
    )

    doc.add_heading("13. Tablas Messiah involucradas (referencia)", level=1)
    doc.add_paragraph(
        "ss_solicitud, ss_solicitud_medicamento, ss_solicitud_atencion, ss_autorizacion (y relacionadas), "
        "af_afiliado, ct_ips, ct_ips_sede, tb_medico_solicitante, tb_cie10, tb_medicamento, "
        "tb_modalidad_servicio_salud, contratos y tarifarios de medicamentos, direccionamiento IPS."
    )

    doc.add_paragraph("")
    p = doc.add_paragraph(
        "Documento generado automáticamente — scripts/generate_manual_ordenamiento_medicamentos.py"
    )
    p.italic = True

    doc.save(out)
    print(f"Escrito: {out}")


if __name__ == "__main__":
    main()
