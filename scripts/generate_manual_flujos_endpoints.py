"""Genera Manual_Flujos_Endpoints_API.docx con el flujo detallado de cada endpoint."""
from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def _add_meta_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    for i, (k, v) in enumerate(rows):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
    doc.add_paragraph("")


def _add_bullet_list(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _add_flow(doc: Document, steps: list[str]) -> None:
    for i, step in enumerate(steps, start=1):
        doc.add_paragraph(f"{i}. {step}", style="List Number")


def _fmt_tables(read: list[str], write: list[str]) -> str:
    parts: list[str] = []
    if read:
        parts.append("Lectura: " + ", ".join(read))
    if write:
        parts.append("Escritura: " + ", ".join(write))
    return "\n".join(parts) if parts else "— (sin acceso directo)"


def _add_endpoint(doc: Document, ep: dict[str, Any]) -> None:
    doc.add_heading(f"{ep['method']} {ep['path']}", level=2)
    doc.add_paragraph(ep["title"])
    _add_meta_table(
        doc,
        [
            ("Grupo OpenAPI", ep["grupo"]),
            ("Autenticación", ep["auth"]),
            ("Autorización adicional", ep["permiso"]),
            ("Content-Type", ep["content_type"]),
            ("Auditoría", ep["auditoria"]),
        ],
    )

    doc.add_heading("Tablas PostgreSQL (ERP Messiah)", level=3)
    doc.add_paragraph(_fmt_tables(ep.get("pg_read", []), ep.get("pg_write", [])))

    doc.add_heading("Tablas SQL Server (OrquestacionDB)", level=3)
    doc.add_paragraph(_fmt_tables(ep.get("sql_read", []), ep.get("sql_write", [])))

    doc.add_heading("Validaciones", level=3)
    _add_bullet_list(doc, ep.get("validaciones", []))

    doc.add_heading("Entrada", level=3)
    doc.add_paragraph(ep["request_desc"])

    doc.add_heading("Salida", level=3)
    doc.add_paragraph(ep["response_desc"])

    doc.add_heading("Flujo detallado", level=3)
    _add_flow(doc, ep["pasos"])

    doc.add_heading("Códigos de respuesta", level=3)
    doc.add_paragraph(ep["codigos"])
    doc.add_paragraph("")


_COMMON_JWT = [
    "Bearer JWT válido (iss, aud, exp, sub, jti).",
    "Usuario activo en orq.usuarios.",
    "IP del cliente dentro de IpPermitida del usuario (si está configurada).",
    "Lista blanca seg.usuario_endpoints para usuarios no administradores (si tienen reglas activas).",
]

_COMMON_LOG = "orq.log_procesos (consumo API, servicio lógico prefijado api:)."


ENDPOINTS: list[dict[str, Any]] = [
    {
        "title": "Inicio de sesión y emisión de JWT.",
        "method": "POST",
        "path": "/auth/login",
        "grupo": "Auth",
        "auth": "Público (sin Bearer).",
        "permiso": "Ninguno.",
        "content_type": "application/json",
        "pg_read": [],
        "pg_write": [],
        "sql_read": ["orq.usuarios", "seg.perfiles", "seg.usuario_perfil"],
        "sql_write": ["orq.log_accesos", "orq.usuarios (ultimo_login)"],
        "validaciones": [
            "Pydantic: username y password obligatorios.",
            "Usuario debe existir y estar activo (orq.usuarios.activo = 1).",
            "Contraseña: hash SHA-256 hex o texto plano transitorio (verify_password_hash).",
            "IP del cliente autorizada según orq.usuarios.IpPermitida (si aplica).",
            "Roles desde seg.perfiles; si usuarios.tipo = 1 se agrega rol ADMINISTRADOR.",
        ],
        "request_desc": (
            "JSON: username, password.\n"
            "Ejemplo: {\"username\": \"api_user\", \"password\": \"***\"}"
        ),
        "response_desc": "200 — TokenResponse: access_token (JWT), expires_in (segundos).",
        "auditoria": "orq.log_accesos (SUCCESS/ERROR, IP, mensaje).",
        "codigos": "200 | 401 credenciales inválidas o usuario inactivo | 403 IP no autorizada | 422",
        "pasos": [
            "Resolver IP del cliente.",
            "Buscar usuario activo en orq.usuarios.",
            "Verificar contraseña; si falla → log ERROR → 401.",
            "Validar lista blanca IP.",
            "Actualizar ultimo_login y registrar log SUCCESS.",
            "Emitir JWT con roles y devolver token.",
        ],
    },
    {
        "title": "Perfil del usuario autenticado y permisos seg.*.",
        "method": "GET",
        "path": "/auth/me",
        "grupo": "Auth",
        "auth": "Bearer JWT.",
        "permiso": "Lista blanca seg.usuario_endpoints si aplica.",
        "content_type": "—",
        "pg_read": [],
        "pg_write": [],
        "sql_read": [
            "orq.usuarios",
            "seg.modulos",
            "seg.acciones",
            "seg.perfiles",
            "seg.permisos",
            "seg.usuario_perfil",
            "seg.usuario_endpoints",
        ],
        "sql_write": [],
        "validaciones": _COMMON_JWT + [
            "Si usuarios.tipo = 1: todos los módulos activos con todas las acciones.",
            "Si no admin: módulos/acciones desde seg.permisos con permitido = true.",
            "503 si no hay permisos SELECT sobre esquema seg.",
        ],
        "request_desc": "Header Authorization: Bearer <token>.",
        "response_desc": "200 — MeResponse: username, es_administrador, modulos[].",
        "auditoria": "No registra orq.log_procesos.",
        "codigos": "200 | 401 | 503 error leyendo seg.*",
        "pasos": [
            "Validar JWT y usuario activo.",
            "Consultar get_me en SqlServerRepository.",
            "Armar listado de módulos y acciones.",
            "Devolver MeResponse.",
        ],
    },
    {
        "title": "Verificación de disponibilidad del servicio.",
        "method": "GET",
        "path": "/health",
        "grupo": "Health",
        "auth": "Bearer JWT.",
        "permiso": "Lista blanca seg.usuario_endpoints si aplica.",
        "content_type": "—",
        "pg_read": [],
        "pg_write": [],
        "sql_read": ["orq.usuarios", "seg.usuario_endpoints"],
        "sql_write": [],
        "validaciones": _COMMON_JWT + [
            "No ejecuta consulta de salud a bases de datos de negocio.",
        ],
        "request_desc": "Sin cuerpo.",
        "response_desc": '200 — {"status": "ok"}.',
        "auditoria": "Sin registro de consumo.",
        "codigos": "200 | 401",
        "pasos": ["Validar JWT.", "Responder status ok."],
    },
    {
        "title": "Historial completo de portabilidad del afiliado (CA1).",
        "method": "POST",
        "path": "/consultas/portabilidad",
        "grupo": "Procesos Afiliado",
        "auth": "Bearer JWT.",
        "permiso": "Lista blanca seg.usuario_endpoints si aplica.",
        "content_type": "application/json",
        "pg_read": [
            "administrativo.af_afiliado",
            "administrativo.af_afiliado_movilidad",
            "administrativo.tb_municipio",
        ],
        "pg_write": [],
        "sql_read": [],
        "sql_write": [_COMMON_LOG],
        "validaciones": _COMMON_JWT + [
            "tipo_identificacion normalizado vía catálogo (CC, TI, RC, etc.).",
            "Afiliado debe existir (404).",
            "CA1: estado_afiliado debe ser Activo (1) o Activo Carnetizado (5); si no → 400.",
        ],
        "request_desc": "JSON: tipo_identificacion, numero_identificacion.",
        "response_desc": "200 — PortabilidadConsultaResponse con historial de movimientos.",
        "auditoria": _COMMON_LOG + " servicio=consulta_portabilidad.",
        "codigos": "200 | 400 afiliado no activo | 404 | 401 | 422 | 500",
        "pasos": [
            "Resolver documento.",
            "Consultar resumen afiliado en PostgreSQL.",
            "Validar estado activo para portabilidad.",
            "Obtener movimientos (fetch_todas_portabilidades_por_afiliado).",
            "Construir respuesta y registrar consumo.",
        ],
    },
    {
        "title": "Listado de PQR del afiliado (POST, filtro opcional por radicado).",
        "method": "POST",
        "path": "/consultas/pqr/por-afiliado",
        "grupo": "Procesos Afiliado",
        "auth": "Bearer JWT.",
        "permiso": "Lista blanca seg.usuario_endpoints si aplica.",
        "content_type": "application/json",
        "pg_read": [
            "administrativo.af_afiliado",
            "administrativo.pqr_peticion_encabezado",
            "administrativo.pqr_tipo_solicitud",
            "administrativo.prb_dependencia",
        ],
        "pg_write": [],
        "sql_read": [],
        "sql_write": [_COMMON_LOG],
        "validaciones": _COMMON_JWT + [
            "Afiliado debe existir (404).",
            "consecutivo_peticion opcional: si se envía, filtra una PQR; si no existe devuelve lista vacía (200).",
        ],
        "request_desc": "JSON: tipo_identificacion, numero_identificacion, consecutivo_peticion (opcional).",
        "response_desc": "200 — PqrAfiliadoListaResponse.",
        "auditoria": _COMMON_LOG + " servicio=consulta_pqr_por_afiliado.",
        "codigos": "200 | 404 | 401 | 422 | 500",
        "pasos": [
            "Verificar afiliado.",
            "Listar PQR o buscar por consecutivo.",
            "Mapear a resumen y registrar consumo.",
        ],
    },
    {
        "title": "Detalle de portabilidad del afiliado (GET).",
        "method": "GET",
        "path": "/consultas/afiliado/portabilidad",
        "grupo": "Procesos Afiliado",
        "auth": "Bearer JWT.",
        "permiso": "Lista blanca seg.usuario_endpoints si aplica.",
        "content_type": "—",
        "pg_read": [
            "administrativo.af_afiliado",
            "administrativo.af_afiliado_movilidad",
            "administrativo.ct_ips",
            "administrativo.tb_zonificacion_ips_encabezado",
            "administrativo.tb_municipio",
        ],
        "pg_write": [],
        "sql_read": [],
        "sql_write": [_COMMON_LOG],
        "validaciones": _COMMON_JWT + [
            "Query: tipo_identificacion, numero_identificacion obligatorios.",
            "Afiliado debe existir (404).",
            "No aplica filtro de estado activo (a diferencia de POST /consultas/portabilidad).",
        ],
        "request_desc": "Query: tipo_identificacion, numero_identificacion.",
        "response_desc": "200 — PortabilidadAfiliadoResponse.",
        "auditoria": _COMMON_LOG + " servicio=consulta_portabilidad_resumen_afiliado.",
        "codigos": "200 | 404 | 401 | 500",
        "pasos": [
            "Verificar afiliado.",
            "Consultar detalle portabilidad.",
            "Devolver ítems y registrar consumo.",
        ],
    },
    {
        "title": "Consulta PQR del afiliado (GET, consecutivo opcional).",
        "method": "GET",
        "path": "/consultas/afiliado/pqr",
        "grupo": "Procesos Afiliado",
        "auth": "Bearer JWT.",
        "permiso": "Lista blanca seg.usuario_endpoints si aplica.",
        "content_type": "—",
        "pg_read": [
            "administrativo.af_afiliado",
            "administrativo.pqr_peticion_encabezado",
            "administrativo.prb_dependencia",
        ],
        "pg_write": [],
        "sql_read": [],
        "sql_write": [_COMMON_LOG],
        "validaciones": _COMMON_JWT + [
            "Afiliado debe existir (404).",
            "consecutivo_pqr opcional en query.",
        ],
        "request_desc": "Query: tipo_identificacion, numero_identificacion; consecutivo_pqr (opcional).",
        "response_desc": "200 — PqrAfiliadoConsultaResponse.",
        "auditoria": _COMMON_LOG + " servicio=consulta_pqr_afiliado_opcional_consecutivo.",
        "codigos": "200 | 404 | 401 | 500",
        "pasos": [
            "Verificar afiliado.",
            "Consultar PQR con filtro opcional.",
            "Registrar consumo.",
        ],
    },
    {
        "title": "Consulta general de datos del afiliado.",
        "method": "POST",
        "path": "/consultas/afiliado",
        "grupo": "Procesos Afiliado",
        "auth": "Bearer JWT.",
        "permiso": "Lista blanca seg.usuario_endpoints si aplica.",
        "content_type": "application/json",
        "pg_read": [
            "administrativo.af_afiliado",
            "administrativo.af_afiliado_complemento",
            "administrativo.tb_municipio",
            "administrativo.ct_ips",
            "administrativo.tb_zonificacion_ips_encabezado",
            "administrativo.af_afiliado_movilidad",
            "administrativo.ct_ips_contrato",
        ],
        "pg_write": [],
        "sql_read": [],
        "sql_write": [_COMMON_LOG],
        "validaciones": _COMMON_JWT + [
            "Afiliado debe existir (404).",
            "Si nit_ips se envía: debe existir en administrativo.ct_ips (404).",
            "Si nit_ips: afiliado debe tener departamento_codigo (400).",
            "Si nit_ips: debe existir contrato IPS para el departamento del afiliado en ct_ips_contrato (404).",
            "Enriquece respuesta con numero_contrato y estado_contrato.",
        ],
        "request_desc": "JSON: tipo_identificacion, numero_identificacion, nit_ips (opcional).",
        "response_desc": "200 — ConsultaAfiliadoResponse.",
        "auditoria": _COMMON_LOG + " servicio=consulta_afiliado.",
        "codigos": "200 | 400 | 404 | 401 | 422 | 500",
        "pasos": [
            "Consultar afiliado.",
            "Si nit_ips: validar IPS y contrato por departamento.",
            "Armar respuesta y registrar consumo.",
        ],
    },
    {
        "title": "Generación de certificado de afiliación en PDF (Base64).",
        "method": "POST",
        "path": "/consultas/certificado-afiliacion",
        "grupo": "Procesos Afiliado",
        "auth": "Bearer JWT.",
        "permiso": "Lista blanca seg.usuario_endpoints si aplica.",
        "content_type": "application/json",
        "pg_read": [
            "administrativo.af_afiliado",
            "administrativo.af_afiliado_complemento",
            "administrativo.ct_ips",
            "administrativo.tb_zonificacion_ips_encabezado",
            "administrativo.af_afiliado_ingreso_contributivo",
            "administrativo.af_aportante",
        ],
        "pg_write": [],
        "sql_read": [],
        "sql_write": [_COMMON_LOG],
        "validaciones": _COMMON_JWT + [
            "Afiliado debe existir (404).",
            "No valida estado activo en código (genera PDF con datos del registro).",
            "Régimen contributivo: incluye relaciones laborales (af_afiliado_ingreso_contributivo, af_aportante).",
            "Plantilla DOCX según régimen; conversión PDF vía LibreOffice (503 si no disponible).",
        ],
        "request_desc": "JSON: tipo_identificacion, numero_identificacion.",
        "response_desc": "200 — CertificadoAfiliacionResponse con archivo_pdf_base64.",
        "auditoria": _COMMON_LOG + " servicio=generacion_certificado_afiliacion.",
        "codigos": "200 | 404 | 503 LibreOffice/plantilla | 500 | 401 | 422",
        "pasos": [
            "Obtener datos afiliado para certificado.",
            "Cargar relaciones laborales si contributivo.",
            "Rellenar plantilla DOCX y convertir a PDF.",
            "Codificar Base64 y registrar consumo.",
        ],
    },
    {
        "title": "Registro de solicitud de actualización de datos desde micrositio.",
        "method": "POST",
        "path": "/afiliados/actualizacion-datos-micrositio",
        "grupo": "Procesos Afiliado",
        "auth": "Bearer JWT.",
        "permiso": "Lista blanca seg.usuario_endpoints si aplica.",
        "content_type": "multipart/form-data",
        "pg_read": [
            "administrativo.af_afiliado",
            "administrativo.tb_tipo_soporte",
        ],
        "pg_write": [
            "administrativo.af_ticket_administrativo_aseguramiento",
            "administrativo.af_ticket_administrativo_aseguramiento_detalle",
            "administrativo.af_ticket_administrativo_aseguramiento_soporte",
        ],
        "sql_read": [],
        "sql_write": [_COMMON_LOG],
        "validaciones": _COMMON_JWT + [
            "barrio, direccion, celular, correo_electronico, observacion obligatorios (422).",
            "Afiliado debe existir (404) y tener edad >= 18 años (400).",
            "telefono/celular: 10 dígitos; correo con formato válido (validar_dato_contacto).",
            "Soporte opcional: tipo en {CN, RC, TI, CC, CE, PA, CD, SC, PE, MS, CERTIFICADO_REGISTRADURIA}.",
            "Soporte: extensión .pdf, .png, .jpg, .jpeg; tamaño máx. TICKET_SUPPORT_MAX_MB; no vacío.",
            "tipo_soporte_documento sin archivo → 400.",
            "Ticket: tipo_proceso=1, origen_solicitud=2 (micrositio).",
            "Archivo soporte guardado en disco (ticket_supports_dir), no en blob de BD.",
        ],
        "request_desc": (
            "Form: tipo_identificacion, numero_identificacion, barrio, direccion, celular, "
            "correo_electronico, observacion; telefono (opc.); soporte_documento (opc.)."
        ),
        "response_desc": "200 — ActualizacionDatosMicrositioResponse con consecutivo_ticket.",
        "auditoria": _COMMON_LOG + " servicio=actualizacion_datos_micrositio.",
        "codigos": "200 | 400 | 404 | 422 | 401 | 500",
        "pasos": [
            "Validar campos y soporte opcional.",
            "Verificar afiliado mayor de edad.",
            "Validar formatos de contacto.",
            "Crear ticket en PostgreSQL.",
            "Registrar consumo.",
        ],
    },
    {
        "title": "Crear cita médica para afiliado activo.",
        "method": "POST",
        "path": "/agendamientos",
        "grupo": "Procesos IPS",
        "auth": "Bearer JWT.",
        "permiso": "Lista blanca seg.usuario_endpoints si aplica.",
        "content_type": "application/json",
        "pg_read": ["administrativo.af_afiliado"],
        "pg_write": [],
        "sql_read": ["orq.agendamiento"],
        "sql_write": ["orq.agendamiento", _COMMON_LOG],
        "validaciones": _COMMON_JWT + [
            "AgendamientoRequest: longitudes de campos y estado 0–5.",
            "hora_cita formato HH:MM[:SS]; fecha_cita + hora_cita >= ahora (400).",
            "Afiliado estado_afiliado IN (1, 5) — activo (400).",
            "Sin cita duplicada: misma sede, profesional, fecha, hora, especialidad, estado (409).",
        ],
        "request_desc": (
            "JSON: sede, tipoDoc, numDoc, tipoDoc_Prof, numDoc_Prof, fecha_cita, hora_cita, "
            "especialidad, estado, usuario_asignacion (opc.), programa (opc.)."
        ),
        "response_desc": "201 — AgendamientoResponse con id_agendamiento.",
        "auditoria": _COMMON_LOG + " servicio=crear_agendamiento.",
        "codigos": "201 | 400 | 409 | 401 | 500",
        "pasos": [
            "Validar fecha/hora.",
            "Verificar afiliado activo en PostgreSQL.",
            "Comprobar conflicto en orq.agendamiento.",
            "Insertar cita y registrar consumo.",
        ],
    },
    {
        "title": "Actualizar estado de una cita existente.",
        "method": "PUT",
        "path": "/agendamientos/{id_agendamiento}",
        "grupo": "Procesos IPS",
        "auth": "Bearer JWT.",
        "permiso": "Lista blanca seg.usuario_endpoints si aplica.",
        "content_type": "application/json",
        "pg_read": [],
        "pg_write": [],
        "sql_read": ["orq.agendamiento"],
        "sql_write": ["orq.agendamiento", _COMMON_LOG],
        "validaciones": _COMMON_JWT + [
            "estado debe ser 0–5 (pendiente, confirmada, cancelada, atendida, no asistió, reprogramada).",
            "id_agendamiento debe existir (404).",
        ],
        "request_desc": "Path: id_agendamiento. JSON: estado.",
        "response_desc": "200 — AgendamientoResponse.",
        "auditoria": _COMMON_LOG + " servicio=editar_agendamiento.",
        "codigos": "200 | 400 | 404 | 401 | 500",
        "pasos": [
            "Actualizar estado en orq.agendamiento.",
            "Si no afecta filas → 404.",
            "Registrar consumo.",
        ],
    },
    {
        "title": "Registrar historia clínica ambulatoria (JSON Version_12.0).",
        "method": "POST",
        "path": "/historia_clinica",
        "grupo": "Procesos IPS",
        "auth": "Bearer JWT.",
        "permiso": "Lista blanca seg.usuario_endpoints si aplica.",
        "content_type": "application/json",
        "pg_read": [
            "administrativo.af_afiliado",
            "administrativo.ct_ips",
            "administrativo.tb_cie10",
            "administrativo.tb_cup",
        ],
        "pg_write": [],
        "sql_read": [],
        "sql_write": [
            "orq.historia_clinica",
            "orq.historia_clinica_actividad",
            _COMMON_LOG,
        ],
        "validaciones": _COMMON_JWT + [
            "Afiliado activo (estado 1 o 5) (404).",
            "NIT prestador debe existir en administrativo.ct_ips (404).",
            "EntidadResponsable.Codigo debe ser CCF033 o CCFC33 (400).",
            "CodigoDiagnosticoPrincipal en tb_cie10 (404).",
            "CodigoCups en tb_cup (404).",
            "Estructura Pydantic HistoriaClinicaRequest (Usuario, Prestador, Cita, Mediciones, Actividades).",
        ],
        "request_desc": "JSON HistoriaClinicaRequest (Version_12.0).",
        "response_desc": "201 — HistoriaClinicaResponse con nueva_historia_id.",
        "auditoria": _COMMON_LOG + " servicio=registrar_historia_clinica.",
        "codigos": "201 | 400 | 404 | 401 | 500",
        "pasos": [
            "Validar afiliado, IPS, CIE-10 y CUPS.",
            "Insertar orq.historia_clinica y actividades.",
            "Registrar consumo.",
        ],
    },
    {
        "title": "Registrar historia clínica desde Bundle FHIR + PDF soporte.",
        "method": "POST",
        "path": "/historia_clinica/hl7",
        "grupo": "Procesos IPS",
        "auth": "Bearer JWT.",
        "permiso": "Lista blanca seg.usuario_endpoints si aplica.",
        "content_type": "application/json",
        "pg_read": [
            "administrativo.af_afiliado",
            "administrativo.ct_ips",
            "administrativo.tb_cie10",
            "administrativo.tb_cup",
        ],
        "pg_write": [],
        "sql_read": [],
        "sql_write": [
            "orq.historia_clinica",
            "orq.historia_clinica_actividad",
            _COMMON_LOG,
        ],
        "validaciones": _COMMON_JWT + [
            "payload.bundle debe ser objeto Bundle FHIR (400).",
            "historia_clinica_pdf_base64 obligatorio; Base64 válido; contenido debe iniciar con %PDF- (400).",
            "Errores de mapeo FHIR → HistoriaClinicaRequest (400).",
            "Luego aplica todas las validaciones de POST /historia_clinica.",
            "PDF no se persiste en BD; solo se registra tamaño en detalle de consumo.",
        ],
        "request_desc": "JSON: bundle (FHIR), historia_clinica_pdf_base64.",
        "response_desc": "201 — HistoriaClinicaResponse (igual que /historia_clinica).",
        "auditoria": _COMMON_LOG + " servicio=registrar_historia_clinica_hl7_fhir.",
        "codigos": "201 | 400 | 404 | 401 | 500",
        "pasos": [
            "Validar bundle y PDF Base64.",
            "Mapear FHIR a HistoriaClinicaRequest.",
            "Delegar en registrar_historia_clinica.",
            "Registrar consumo FHIR.",
        ],
    },
    {
        "title": "Paso 1 — Solicitud y autorización emitida de medicamentos por orden médica IPS.",
        "method": "POST",
        "path": "/afiliados/autorizacion-orden-medica-ips",
        "grupo": "Procesos IPS",
        "auth": "Bearer JWT.",
        "permiso": "orq.usuarios.autoriza_med = 1 (require_autoriza_med).",
        "content_type": "multipart/form-data",
        "pg_read": [
            "administrativo.af_afiliado",
            "administrativo.af_afiliado_complemento",
            "administrativo.tb_municipio",
            "administrativo.ct_ips",
            "administrativo.ct_ips_sede",
            "administrativo.tb_medicamento",
            "administrativo.tb_unidad_medida",
            "administrativo.ct_ips_contrato",
            "administrativo.ct_ips_contrato_cobertura",
            "administrativo.tb_tarifario_propio_detalle",
            "administrativo.tb_direccionamiento_autorizacion",
            "administrativo.tb_direccionamiento_autorizacion_medicamento",
            "administrativo.tb_medico_solicitante",
            "administrativo.tb_especialidad",
            "administrativo.tb_cie10",
            "administrativo.tb_modalidad_servicio_salud",
            "administrativo.ss_autorizacion",
            "administrativo.ss_autorizacion_medicamento",
            "administrativo.tb_tipo_soporte",
            "administrativo.ss_solicitud",
        ],
        "pg_write": [
            "administrativo.ct_ips_ss_solicitud",
            "administrativo.tb_parametro",
            "administrativo.ss_solicitud",
            "administrativo.ss_solicitud_atencion",
            "administrativo.ss_solicitud_medicamento",
            "administrativo.ss_autorizacion",
            "administrativo.ss_autorizacion_medicamento",
            "administrativo.ct_ips_ss_solicitud_autorizacion",
            "administrativo.ss_solicitud_soporte",
        ],
        "sql_read": ["orq.usuarios (autoriza_med)"],
        "sql_write": [_COMMON_LOG],
        "validaciones": _COMMON_JWT + [
            "orq.usuarios.autoriza_med = true (403).",
            "medicamentos_json: array JSON válido de objetos con CUM, cantidad, días (400).",
            "Fechas solicitud no futuras (fecha_solicitud_proceso, fecha_solicitud_medico).",
            "observacion obligatoria; diagnóstico principal CIE-10; registro_profesional en tb_medico_solicitante.",
            "Sin CUMs duplicados en el listado.",
            "Afiliado estado_afiliado = 1 (Activo) (400).",
            "NIT IPS solicitante y direccionamiento deben resolverse en ct_ips (404).",
            "email obligatorio (formulario o sede/afiliado/IPS); telefono y celular 10 dígitos.",
            "soporte_orden_medica PDF obligatorio; validación %PDF- y tamaño máximo.",
            "Origen: Orden médica; atención Enfermedad General; modalidad Ambulatorios.",
            "Por medicamento: sw_activo, tarifario municipio, contrato vigente, direccionamiento IPS.",
            "Al menos un medicamento debe autorizarse; si ninguno cumple no se crea ss_solicitud (HTTP 400).",
            "Autorización estado EMITIDA con PIN; autorizacion_activa=false hasta activación.",
        ],
        "request_desc": (
            "multipart/form-data: documento afiliado, diagnósticos, NIT IPS, fechas, "
            "medicamentos_json, soporte_orden_medica PDF, contacto."
        ),
        "response_desc": "200 — AutorizacionOrdenMedicaIpsResponse con PIN y listados de medicamentos.",
        "auditoria": _COMMON_LOG + " servicio=autorizacion_orden_medica_ips.",
        "codigos": "200 | 400 | 403 | 404 | 401 | 422 | 500",
        "pasos": [
            "Validar autoriza_med y parsear formulario.",
            "Validar afiliado, médico, IPS y medicamentos.",
            "Crear ss_solicitud y evaluar reglas por CUM.",
            "Emitir ss_autorizacion con PIN si aplica.",
            "Opcional: PDF JasperStarter.",
            "Registrar traza en orq.log_procesos.",
        ],
    },
    {
        "title": "Paso 2 — Activación y confirmación de autorización IPS con PIN.",
        "method": "POST",
        "path": "/afiliados/autorizacion-orden-medica-ips/activar",
        "grupo": "Procesos IPS",
        "auth": "Bearer JWT.",
        "permiso": "orq.usuarios.autoriza_med = 1.",
        "content_type": "multipart/form-data",
        "pg_read": [
            "administrativo.ct_ips",
            "administrativo.tb_municipio",
            "administrativo.ss_autorizacion",
            "administrativo.ss_autorizacion_medicamento",
            "administrativo.ss_solicitud",
            "administrativo.af_afiliado",
        ],
        "pg_write": [
            "administrativo.ss_autorizacion (fecha_activacion, fecha_real_prestacion_servicio, estado_trazabilidad, url_activacion, sw_activo)",
            "administrativo.ss_autorizacion_medicamento (fecha_programacion, fecha_prestacion_servicio)",
        ],
        "sql_read": ["orq.usuarios (autoriza_med)"],
        "sql_write": [_COMMON_LOG],
        "validaciones": _COMMON_JWT + [
            "orq.usuarios.autoriza_med = true (403).",
            "pin_activacion obligatorio (comparación sin distinguir mayúsculas).",
            "NIT IPS direccionamiento debe resolverse (404).",
            "Autorización por documento + PIN + IPS: emitida, activada sin confirmar o ya completada (404 si no existe).",
            "fecha_programacion: anterior a fin vigencia y no anterior a fecha grabado.",
            "fecha_prestacion: no futura, no anterior a fecha grabado ni a fecha_programacion.",
            "confirmar_prestacion default true; soporte_confirmacion PDF opcional.",
            "Idempotente si ya activada o confirmada (200).",
        ],
        "request_desc": (
            "Form: tipo_identificacion, numero_identificacion, pin_activacion, nit_ips_direccionamiento, "
            "fecha_programacion, fecha_prestacion, confirmar_prestacion, soporte_confirmacion (PDF opcional)."
        ),
        "response_desc": (
            "200 — ActivacionOrdenMedicaIpsResponse con prestacion_confirmada, PDF prestado (PRESTADO=1) "
            "y soporte_confirmacion_registrado_messiah cuando aplica."
        ),
        "auditoria": _COMMON_LOG + " servicio=activacion_autorizacion_orden_medica_ips.",
        "codigos": "200 | 400 | 403 | 404 | 409 | 401 | 500",
        "pasos": [
            "Validar autoriza_med.",
            "Buscar autorización por documento, PIN e IPS.",
            "Programar medicamentos y activar (ACTIVADA=2) si pendiente.",
            "Confirmar prestación (COMPLETADA=3) si confirmar_prestacion=true.",
            "Guardar soporte opcional en repositorio Messiah autorizacion_confirmada/.",
            "PDF Jasper reAutorizacion con PRESTADO=1.",
            "Registrar consumo.",
        ],
    },
    {
        "title": "Ajuste decimal de facturas (proceso financiero batch).",
        "method": "POST",
        "path": "/execute/facturas_decimales",
        "grupo": "Procesos Financieros",
        "auth": "Bearer JWT.",
        "permiso": "Módulo «Facturas decimales», acción EJECUTAR (seg.permisos) o admin.",
        "content_type": "— (sin cuerpo)",
        "pg_read": [
            "administrativo.sc_saldo_encabezado",
            "administrativo.sc_saldo_detalle",
            "administrativo.sc_cuenta",
            "administrativo.sc_tercero",
            "administrativo.sc_factura_encabezado",
            "administrativo.ct_ips",
        ],
        "pg_write": [],
        "sql_read": [
            "orq.usuarios",
            "seg.permisos",
            "seg.perfiles",
            "seg.usuario_perfil",
            "seg.modulos",
            "seg.acciones",
            "orq.resultados_procesos",
        ],
        "sql_write": ["orq.resultados_procesos", "orq.log_procesos"],
        "validaciones": _COMMON_JWT + [
            "Permiso módulo «Facturas decimales» + acción «EJECUTAR» (403); admin tipo=1 omite.",
            "Consulta candidatos: delta saldo <> 0, saldo_factura > 0, cuenta clase_b=1.",
            "Tercero NIT coincide con ct_ips.nit; fechas factura 2022–2030.",
            "Omite referencias ya registradas en orq.resultados_procesos (tipo_proceso=facturas_decimales).",
        ],
        "request_desc": "Sin cuerpo. Solo Authorization Bearer.",
        "response_desc": "200 — ApiExecutionResponse: total, success, errors, detail.",
        "auditoria": "orq.log_procesos servicio=facturas_decimales.",
        "codigos": "200 | 403 | 401 | 500 | 503",
        "pasos": [
            "Validar permiso Facturas decimales.",
            "Consultar candidatos en PostgreSQL.",
            "Persistir en orq.resultados_procesos evitando duplicados.",
            "Registrar log y devolver estadísticas.",
        ],
    },
    {
        "title": "Actualización unificada de saldo y valor por aplicar de facturas.",
        "method": "POST",
        "path": "/execute/update_saldo_y_valor_factura",
        "grupo": "Procesos Financieros",
        "auth": "Bearer JWT.",
        "permiso": "Módulo «Actualizacion saldo y valor», acción EJECUTAR o admin.",
        "content_type": "— (sin cuerpo)",
        "pg_read": [
            "administrativo.sc_saldo_encabezado",
            "administrativo.sc_saldo_detalle",
            "administrativo.sc_cuenta",
            "administrativo.sc_tercero",
            "administrativo.sc_factura_encabezado",
            "administrativo.ct_ips",
        ],
        "pg_write": [
            "administrativo.sc_factura_encabezado (saldo_factura)",
            "administrativo.sc_factura_detalle_valor (valor_por_aplicar)",
        ],
        "sql_read": [
            "orq.usuarios",
            "seg.permisos",
            "seg.perfiles",
            "seg.usuario_perfil",
            "seg.modulos",
            "seg.acciones",
        ],
        "sql_write": ["orq.log_procesos"],
        "validaciones": _COMMON_JWT + [
            "Permiso módulo «Actualizacion saldo y valor» + acción «EJECUTAR» (403).",
            "Misma consulta de candidatos que facturas_decimales.",
            "Aplica int(valor) por consecutivo_factura en lotes de 500.",
            "Commit al éxito; rollback ante error.",
        ],
        "request_desc": "Sin cuerpo.",
        "response_desc": "200 — ApiExecutionResponse.",
        "auditoria": "orq.log_procesos servicio=update_saldo_y_valor_factura.",
        "codigos": "200 | 403 | 401 | 500 | 503",
        "pasos": [
            "Validar permiso.",
            "Obtener payload desde PostgreSQL.",
            "Actualizar saldo_factura y valor_por_aplicar por lote.",
            "Registrar log y devolver respuesta.",
        ],
    },
    {
        "title": "Alias legacy del ajuste decimal (oculto en OpenAPI).",
        "method": "POST",
        "path": "/execute/service1",
        "grupo": "Ejecución (legacy)",
        "auth": "Bearer JWT.",
        "permiso": "Igual que /execute/facturas_decimales.",
        "content_type": "—",
        "pg_read": [
            "administrativo.sc_saldo_encabezado",
            "administrativo.sc_saldo_detalle",
            "administrativo.sc_cuenta",
            "administrativo.sc_tercero",
            "administrativo.sc_factura_encabezado",
            "administrativo.ct_ips",
        ],
        "pg_write": [],
        "sql_read": [
            "orq.usuarios",
            "seg.permisos",
            "seg.perfiles",
            "seg.usuario_perfil",
            "orq.resultados_procesos",
        ],
        "sql_write": ["orq.resultados_procesos", "orq.log_procesos"],
        "validaciones": [
            "Idénticas a POST /execute/facturas_decimales.",
        ],
        "request_desc": "Sin cuerpo.",
        "response_desc": "200 — ApiExecutionResponse.",
        "auditoria": "Igual que facturas_decimales.",
        "codigos": "200 | 403 | 401 | 500",
        "pasos": [
            "Validar permiso Facturas decimales.",
            "Ejecutar OrchestratorService.run_service1.",
            "Devolver respuesta.",
        ],
    },
]


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    out = root / "Manual_Flujos_Endpoints_API.docx"

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Flujos detallados — API ORQUESTADORDB", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "EPS Familiar de Colombia — Backend FastAPI que orquesta operaciones entre "
        "PostgreSQL (ERP Messiah, esquema administrativo) y SQL Server (OrquestacionDB)."
    )
    doc.add_paragraph(
        "Por cada endpoint se documentan: tablas PostgreSQL y SQL Server (lectura/escritura), "
        "validaciones de negocio y seguridad, flujo paso a paso. "
        "Radicación SIIFA: scripts externos (sin endpoint API)."
    )

    doc.add_heading("1. Convenciones transversales", level=1)
    doc.add_paragraph(
        "Autenticación JWT en todas las rutas excepto POST /auth/login.\n\n"
        "Lista blanca: seg.usuario_endpoints (método + ruta) para usuarios no administradores.\n\n"
        "Auditoría API: orq.log_procesos con servicio lógico, usuario, documento, resultado y detalle JSON.\n\n"
        "Esquema PostgreSQL principal: administrativo.* (Messiah ERP)."
    )

    doc.add_heading("2. Índice de endpoints", level=1)
    idx_table = doc.add_table(rows=1, cols=4)
    hdr = idx_table.rows[0].cells
    hdr[0].text = "Grupo"
    hdr[1].text = "Método"
    hdr[2].text = "Ruta"
    hdr[3].text = "Resumen"
    for ep in ENDPOINTS:
        row = idx_table.add_row().cells
        row[0].text = ep["grupo"]
        row[1].text = ep["method"]
        row[2].text = ep["path"]
        row[3].text = ep["title"]
    doc.add_paragraph("")

    doc.add_heading("3. Procesos SIIFA", level=1)
    doc.add_paragraph(
        "Sin endpoints API vigentes. Radicación vía:\n"
        "• export/siifa_radicacion_sin_seguimiento_api.py\n"
        "• scripts/siifa_radicacion_sync.py\n"
        "Tablas típicas: dbo.SIIFA_* (SQL Server), administrativo.rips_af / rips_resumen (PostgreSQL)."
    )

    current_group = ""
    doc.add_heading("4. Flujos por endpoint", level=1)
    for ep in ENDPOINTS:
        if ep["grupo"] != current_group:
            current_group = ep["grupo"]
            doc.add_heading(current_group, level=2)
        _add_endpoint(doc, ep)

    doc.add_heading("5. Códigos HTTP habituales", level=1)
    doc.add_paragraph(
        "200/201 — Éxito | 400 — Validación negocio | 401 — Token/login | 403 — Permisos/ACL\n"
        "404 — No encontrado | 409 — Conflicto | 422 — Esquema entrada | 500 — Error interno | 503 — Servicio no disponible"
    )

    doc.add_paragraph("")
    p = doc.add_paragraph(
        "Generado por scripts/generate_manual_flujos_endpoints.py — ORQUESTADORDB."
    )
    p.italic = True

    try:
        doc.save(out)
    except PermissionError:
        out = root / "Manual_Flujos_Endpoints_API_v2.docx"
        doc.save(out)
    print(f"Escrito: {out}")


if __name__ == "__main__":
    main()
