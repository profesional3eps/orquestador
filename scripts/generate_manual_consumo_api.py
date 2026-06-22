"""Genera Manual_Consumo_API_ORQUESTADORDB.docx en la raíz del proyecto."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    out = root / "Manual_Consumo_API_ORQUESTADORDB.docx"

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Manual de consumo — API ORQUESTADORDB", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "Backend FastAPI que orquesta operaciones entre PostgreSQL y SQL Server (OrquestacionDB). "
        "Este documento describe autenticación, autorización y uso de los endpoints REST."
    )

    doc.add_heading("1. Alcance y URL base", level=1)
    doc.add_paragraph(
        "La URL base depende del despliegue (por ejemplo http://localhost:8000 en desarrollo). "
        "Todos los paths de este manual son relativos a esa base."
    )
    doc.add_paragraph(
        "Documentación interactiva: GET /docs (Swagger UI), GET /redoc (ReDoc), GET /openapi.json (OpenAPI 3)."
    )

    doc.add_heading("2. Autenticación JWT", level=1)
    doc.add_paragraph("2.1 Login (público, sin Bearer)")
    doc.add_paragraph(
        "Método y ruta: POST /auth/login\n"
        "Content-Type: application/json\n"
        "Cuerpo ejemplo: {\"username\": \"<usuario>\", \"password\": \"<contraseña>\"}"
    )
    doc.add_paragraph(
        "Respuesta 200: objeto con access_token (JWT) y expires_in (segundos). "
        "Respuesta 401: credenciales inválidas o usuario inactivo."
    )
    doc.add_paragraph(
        "Las credenciales se validan contra orq.usuarios (activo=1). "
        "La contraseña puede estar almacenada como hash SHA-256 (hex 64 caracteres) o, "
        "en modo transitorio, en texto plano según implementación del hash en el servidor."
    )

    doc.add_paragraph("2.2 Peticiones protegidas")
    doc.add_paragraph(
        "Incluir en cada solicitud (excepto POST /auth/login):\n"
        "Authorization: Bearer <access_token>"
    )
    doc.add_paragraph(
        "Si el token falta, es inválido o expiró, la API responde 401."
    )

    doc.add_heading("3. Autorización", level=1)
    doc.add_heading("3.1 Lista blanca por endpoint (seg.usuario_endpoints)", level=2)
    doc.add_paragraph(
        "Para usuarios que no son administradores (orq.usuarios.tipo = 0), si existen filas "
        "activas en seg.usuario_endpoints para ese id_usuario, solo se permiten las rutas "
        "definidas allí, comparando método HTTP (ej. POST, GET, PUT) y la ruta del endpoint "
        "tal como la registra FastAPI (incluye parámetros de ruta en forma de plantilla, "
        "por ejemplo /agendamientos/{id_agendamiento})."
    )
    doc.add_paragraph(
        "Si el usuario no tiene ninguna regla en seg.usuario_endpoints, se mantiene el "
        "comportamiento anterior: no se bloquea por lista blanca (salvo otros controles)."
    )
    doc.add_paragraph(
        "Los administradores (orq.usuarios.tipo = 1) omiten la lista blanca por endpoint."
    )
    doc.add_paragraph(
        "Script de ejemplo para crear la tabla: sql/seg_usuario_endpoints.sql"
    )

    doc.add_heading("3.2 Permisos por módulo y acción (seg.permisos)", level=2)
    doc.add_paragraph(
        "Algunos endpoints exigen además permiso en seg.* (perfiles, permisos, módulos, acciones) "
        "mediante seg.permisos y seg.usuario_perfil, o bien usuario administrador (tipo = 1)."
    )
    doc.add_paragraph(
        "• POST /execute/facturas_decimales y POST /execute/service1: EJECUTAR sobre «Facturas decimales».\n"
        "• POST /execute/update_saldo_y_valor_factura: EJECUTAR sobre «Actualizacion saldo y valor»."
    )

    doc.add_heading("3.3 Perfil del usuario autenticado", level=2)
    doc.add_paragraph(
        "GET /auth/me devuelve el nombre de usuario, si es administrador y el listado de módulos "
        "con acciones permitidas. Requiere Bearer y está sujeto a la lista blanca por endpoint si aplica."
    )

    doc.add_heading("4. Catálogo de endpoints", level=1)
    doc.add_paragraph(
        "La tabla resume método, ruta, si requiere JWT y notas. Los cuerpos y respuestas detallados "
        "están en /openapi.json y en /docs."
    )

    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Método"
    hdr[1].text = "Ruta"
    hdr[2].text = "JWT"
    hdr[3].text = "Notas"

    rows = [
        ("POST", "/auth/login", "No", "Emite token"),
        ("GET", "/auth/me", "Sí", "Perfil y permisos seg.*"),
        ("GET", "/health", "Sí", 'Respuesta {"status": "ok"}'),
        ("POST", "/consultas/portabilidad", "Sí", "JSON: tipo_identificacion, numero_identificacion; afiliado activo"),
        ("POST", "/consultas/pqr/por-afiliado", "Sí", "JSON: consecutivo_peticion opcional"),
        ("GET", "/consultas/afiliado/portabilidad", "Sí", "Query: tipo_identificacion, numero_identificacion"),
        ("GET", "/consultas/afiliado/pqr", "Sí", "Query opcional consecutivo_pqr"),
        ("POST", "/consultas/afiliado", "Sí", "Consulta afiliado por documento; nit_ips opcional"),
        ("POST", "/consultas/certificado-afiliacion", "Sí", "PDF Base64; afiliación vigente"),
        ("POST", "/afiliados/autorizacion-orden-medica-ips/activar", "Sí", "multipart/form-data; activación + confirmación prestación, PDF prestado"),
        ("POST", "/agendamientos", "Sí", "Crear cita; valida afiliado activo"),
        ("PUT", "/agendamientos/{id_agendamiento}", "Sí", "Actualiza estado del agendamiento"),
        ("POST", "/historia_clinica", "Sí", "JSON Version_12.0; orq.historia_clinica*"),
        ("POST", "/afiliados/actualizacion-datos-micrositio", "Sí", "multipart/form-data"),
        ("POST", "/afiliados/autorizacion-orden-medica-ips", "Sí", "multipart/form-data; orden médica"),
        ("POST", "/execute/facturas_decimales", "Sí", "Permiso Facturas decimales"),
        ("POST", "/execute/update_saldo_y_valor_factura", "Sí", "Permiso Actualizacion saldo y valor"),
        ("POST", "/execute/service1", "Sí", "Alias legacy; mismo permiso que facturas_decimales; oculto en OpenAPI"),
    ]
    for m, path, jwt, note in rows:
        row = table.add_row().cells
        row[0].text = m
        row[1].text = path
        row[2].text = jwt
        row[3].text = note

    doc.add_heading("5. Códigos de respuesta habituales", level=1)
    doc.add_paragraph(
        "• 200/201: éxito según operación.\n"
        "• 400: validación de negocio o datos.\n"
        "• 401: token ausente/inválido o login fallido.\n"
        "• 403: sin permiso de módulo/acción o sin autorización en seg.usuario_endpoints.\n"
        "• 404: recurso no encontrado (ej. afiliado, agendamiento).\n"
        "• 409: conflicto (ej. cita duplicada en agendamiento).\n"
        "• 422: error de validación de esquema (Pydantic).\n"
        "• 500: error interno.\n"
        "• 502/503: fallas de servicios externos (ej. SIIFA) o, en algunos casos, "
        "imposibilidad de leer seg.* para permisos (503 con mensaje operativo)."
    )

    doc.add_heading("6. Trazas y auditoría", level=1)
    doc.add_paragraph(
        "Los intentos de login se registran en orq.log_accesos. "
        "Muchas operaciones de negocio registran consumo en orq.log_procesos con detalle JSON "
        "(servicio lógico, usuario, documento cuando aplica, resultado, http_status)."
    )

    doc.add_heading("7. Colección Postman", level=1)
    doc.add_paragraph(
        "En el repositorio: OrquestadorDB.postman_collection.json. "
        "Use variables baseUrl, apiLoginUsername, apiLoginPassword y token "
        "(el login guarda el token automáticamente con el script de prueba de la petición)."
    )

    doc.add_heading("8. Configuración del servidor (referencia)", level=1)
    doc.add_paragraph(
        "Variables típicas en .env: POSTGRES_URL, SQLSERVER_URL, JWT_SECRET, JWT_ALGORITHM, "
        "JWT_EXPIRE_MINUTES, URLs y credenciales SIIFA cuando se use sincronización, etc. "
        "No incluya secretos en clientes; el consumidor solo necesita URL pública, usuario API y contraseña."
    )

    doc.add_paragraph("")
    p = doc.add_paragraph("Documento generado automáticamente a partir del código del proyecto ORQUESTADORDB.")
    p.italic = True

    doc.save(out)
    print(f"Escrito: {out}")


if __name__ == "__main__":
    main()
