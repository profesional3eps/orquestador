"""Genera Cronograma_Desarrollo_API_Ordenamiento_Medicamentos.docx en la raíz del proyecto."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row_data in rows:
        row = table.add_row().cells
        for i, cell in enumerate(row_data):
            row[i].text = cell
    doc.add_paragraph("")


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    out = root / "Cronograma_Desarrollo_API_Ordenamiento_Medicamentos.docx"

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading(
        "Cronograma de desarrollo — API ORQUESTADORDB",
        0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Estimación de actividades y horas de desarrollo normal para los endpoints "
        "POST /auth/login, POST /consultas/afiliado, "
        "POST /afiliados/autorizacion-orden-medica-ips (paso 1) y "
        "POST /afiliados/autorizacion-orden-medica-ips/activar (paso 2)."
    )
    doc.add_paragraph(
        "Base: alcance implementado en el repositorio ORQUESTADORDB e integración con Messiah "
        "(PostgreSQL administrativo) y orquestador (SQL Server). "
        "Horas orientativas para un desarrollador con experiencia en Python/FastAPI y "
        "conocimiento funcional del dominio EPS (6–8 h/día productivas)."
    )

    doc.add_heading("1. Resumen ejecutivo", level=1)
    _add_table(
        doc,
        ["Bloque", "Horas mín.", "Horas máx.", "Días (8 h)"],
        [
            ["Infraestructura compartida (API, BD, seguridad base)", "48", "64", "6 – 8"],
            ["POST /auth/login", "16", "24", "2 – 3"],
            ["POST /consultas/afiliado", "24", "32", "3 – 4"],
            ["POST /afiliados/autorizacion-orden-medica-ips (paso 1)", "88", "112", "11 – 14"],
            ["POST /afiliados/autorizacion-orden-medica-ips/activar (paso 2)", "104", "136", "13 – 17"],
            ["Pruebas integrales, despliegue y documentación", "40", "56", "5 – 7"],
            ["TOTAL PROYECTO", "320", "424", "40 – 53"],
        ],
    )
    doc.add_paragraph(
        "Rango recomendado para planificación: 360 horas (~45 días hábiles) con un desarrollador, "
        "o 22–24 días con dos desarrolladores en paralelo (uno en flujo Messiah/paso 2, otro en auth/consulta/paso 1)."
    )

    doc.add_heading("2. Infraestructura compartida (prerrequisito)", level=1)
    doc.add_paragraph(
        "Componentes transversales que habilitan los cuatro endpoints. "
        "Se planifican al inicio del proyecto."
    )
    _add_table(
        doc,
        ["#", "Actividad", "Entregable", "Horas"],
        [
            ["2.1", "Análisis técnico ORQUESTADORDB + Messiah (tablas ss_*, ct_ips, sc_saldo)", "Acta de alcance / mapa de tablas", "8"],
            ["2.2", "Proyecto FastAPI: estructura app/, config, settings (.env), logging", "Esqueleto ejecutable", "8"],
            ["2.3", "Conexión dual BD: PostgreSQL (Messiah) + SQL Server (orq.usuarios, auditoría)", "database.py, sesiones Depends", "12"],
            ["2.4", "Docker / docker-compose (Python, variables, healthcheck)", "Imagen desplegable", "8"],
            ["2.5", "OpenAPI / Swagger, DTOs base, manejo errores HTTP", "GET /docs operativo", "8"],
            ["2.6", "Módulo seguridad JWT (Bearer), permisos autoriza_med", "core/security.py, permissions.py", "12"],
            ["2.7", "Registro de consumo API (auditoría SQL Server)", "Traza por servicio/usuario/IP", "6"],
            ["2.8", "Resolución tipo documento, zona horaria Bogotá", "tipo_identificacion.py, zona_horaria.py", "4"],
            ["", "Subtotal infraestructura", "", "48 – 64"],
        ],
    )

    doc.add_heading("3. POST /auth/login", level=1)
    doc.add_paragraph(
        "Autenticación de usuarios del orquestador (orq.usuarios). Emite JWT para consumir rutas protegidas. "
        "Única ruta pública del flujo."
    )
    _add_table(
        doc,
        ["#", "Actividad", "Detalle técnico", "Horas"],
        [
            ["3.1", "Análisis requisitos login", "Usuario activo, hash contraseña, roles, expiración token", "2"],
            ["3.2", "DTO LoginRequest / TokenResponse", "JSON application/json", "2"],
            ["3.3", "Repositorio SQL Server usuarios", "get_active_user_by_username, roles, last_login", "4"],
            ["3.4", "Verificación contraseña (hash)", "verify_password_hash", "2"],
            ["3.5", "Emisión JWT con claims compatibles", "sub, roles, iss, aud, exp; PyJWT", "4"],
            ["3.6", "Whitelist IP por usuario (ip_permitida)", "403 si IP no autorizada; X-Forwarded-For", "4"],
            ["3.7", "Log de acceso (éxito/error)", "AccessLogDTO en SQL Server", "2"],
            ["3.8", "Pruebas unitarias e integración Postman", "Casos 401, 403, 200", "4"],
            ["", "Subtotal /auth/login", "", "16 – 24"],
        ],
    )

    doc.add_heading("4. POST /consultas/afiliado", level=1)
    doc.add_paragraph(
        "Consulta datos generales del afiliado en Messiah por tipo y número de documento. "
        "Opcionalmente valida contrato IPS por NIT y departamento del afiliado."
    )
    _add_table(
        doc,
        ["#", "Actividad", "Detalle técnico", "Horas"],
        [
            ["4.1", "Análisis consulta afiliado Messiah", "Campos requeridos por consumidor IPS", "4"],
            ["4.2", "SQL consulta afiliado (PostgreSQL)", "JOIN afiliado, municipio, régimen, IPS, estado", "8"],
            ["4.3", "DTO ConsultaAfiliadoRequest / Response", "Mapeo campos y normalización", "4"],
            ["4.4", "Validación nit_ips opcional", "ct_ips + contrato por departamento", "6"],
            ["4.5", "Endpoint routes + Bearer obligatorio", "404 afiliado, 400 sin departamento", "4"],
            ["4.6", "Registro consumo API y pruebas", "Postman + datos reales Comfasucre", "6"],
            ["", "Subtotal /consultas/afiliado", "", "24 – 32"],
        ],
    )

    doc.add_heading("5. POST /afiliados/autorizacion-orden-medica-ips (paso 1)", level=1)
    doc.add_paragraph(
        "Solicitud de medicamentos por orden médica IPS. Evalúa reglas antes de persistir; "
        "crea ss_solicitud y ct_ips_ss_solicitud sin ss_autorizacion. Multipart con PDF obligatorio."
    )
    _add_table(
        doc,
        ["#", "Actividad", "Detalle técnico", "Horas"],
        [
            ["5.1", "Análisis flujo Messiah authorization_request_ips", "Réplica reglas Enfermedad General / Ambulatorios", "12"],
            ["5.2", "Servicio direccionamiento (evaluación medicamentos)", "Contratos, tarifario CUM, municipio afiliado", "24"],
            ["5.3", "Validaciones medicamento (Messiah)", "Cantidades, vigencia, concepto nota técnica", "16"],
            ["5.4", "Repositorio Messiah: crear solicitud", "ss_solicitud, ss_solicitud_medicamento, secuencias", "20"],
            ["5.5", "ct_ips_ss_solicitud + campos UI Messiah", "sw_terminada_hospitalaria, servicio Medicamentos", "8"],
            ["5.6", "Validación CIE-10, médico, IPS solicitante/direccionamiento", "postgres_repository + formularios", "8"],
            ["5.7", "Soporte orden médica PDF (multipart)", "Validación tamaño/tipo, almacenamiento Messiah", "12"],
            ["5.8", "Servicio procesar_autorizacion_medicamentos (paso 1)", "Transacción, lock advisory, idempotencia", "16"],
            ["5.9", "DTO respuesta extenso + OpenAPI multipart", "medicamentos autorizados / no autorizados", "8"],
            ["5.10", "Permiso require_autoriza_med + trazabilidad auditoría", "403 sin permiso", "4"],
            ["5.11", "Pruebas integración paso 1", "TOTAL, PARCIAL, 400 sin medicamentos", "12"],
            ["", "Subtotal paso 1", "", "88 – 112"],
        ],
    )

    doc.add_heading("6. POST /afiliados/autorizacion-orden-medica-ips/activar (paso 2)", level=1)
    doc.add_paragraph(
        "Emisión ss_autorizacion, activación, confirmación de prestación, contabilización NC-AT "
        "(preferencias 288/390), PDF Jasper en base64 y soporte confirmación."
    )
    _add_table(
        doc,
        ["#", "Actividad", "Detalle técnico", "Horas"],
        [
            ["6.1", "Análisis SsAutorizacionActivacion Messiah", "Estados trazabilidad 2 (ACTIVADA) y 3 (COMPLETADA)", "12"],
            ["6.2", "Emisión autorización desde solicitud", "ss_autorizacion, ss_autorizacion_medicamento, PIN", "20"],
            ["6.3", "Activación y programación fechas", "fecha_activacion, fecha_programacion medicamentos", "12"],
            ["6.4", "Confirmación prestación + soporte PDF", "url_activacion, ss_solicitud auditoría", "12"],
            ["6.5", "Preferencias tb_preferencia 288/390", "messiah_preferencia_service, cuándo contabilizar", "8"],
            ["6.6", "Contabilización sc_saldo_encabezado/detalle", "NC-AT, sc_contabiliza, validación tipo_documento", "24"],
            ["6.7", "Concurrencia (pg_advisory_xact_lock)", "Alta concurrencia sin colisión de saldos", "8"],
            ["6.8", "PDF Jasper (reAutorizacion.jasper)", "Java 8, JasperStarter, .jrxml, parámetro AUTORIZACION", "32"],
            ["6.9", "Caché compilación Jasper + optimización latencia", "JASPER_CACHE_DIR, reducción ~19 s a ~10 s", "8"],
            ["6.10", "Idempotencia (ya activada / ya confirmada)", "Reintentos seguros paso 2", "8"],
            ["6.11", "Servicio activar_autorizacion_orden_medica_ips", "Orquestación completa paso 2", "16"],
            ["6.12", "Pruebas alineación Messiah UI + contabilidad", "Consulta autorización, saldo NC-AT", "16"],
            ["", "Subtotal paso 2", "", "104 – 136"],
        ],
    )

    doc.add_heading("7. Pruebas, despliegue y documentación", level=1)
    _add_table(
        doc,
        ["#", "Actividad", "Entregable", "Horas"],
        [
            ["7.1", "Colección Postman + variables entorno", "Flujo login → consulta → paso 1 → paso 2", "8"],
            ["7.2", "Manual consumo API (.docx)", "Manual_Consumo_Ordenamiento_Medicamentos.docx", "12"],
            ["7.3", "Pruebas carga / concurrencia paso 2", "Validar locks y tiempos PDF", "8"],
            ["7.4", "Dockerfile Java 8 + JasperStarter", "PDF en contenedor productivo", "8"],
            ["7.5", "Despliegue QA / producción + smoke tests", "docker compose up, verificación BD", "8"],
            ["7.6", "Scripts reparación y soporte operación", "reparar_saldo, reparar solicitud UI", "4"],
            ["", "Subtotal cierre", "", "40 – 56"],
        ],
    )

    doc.add_heading("8. Secuencia sugerida para el cronograma", level=1)
    doc.add_paragraph(
        "Fase 1 (semanas 1–2): Infraestructura + /auth/login + /consultas/afiliado.\n"
        "Fase 2 (semanas 3–5): Paso 1 — direccionamiento, solicitud, validaciones, soporte PDF.\n"
        "Fase 3 (semanas 6–8): Paso 2 — activación, confirmación, contabilización.\n"
        "Fase 4 (semanas 9–10): PDF Jasper, optimización, pruebas Messiah UI, documentación y despliegue."
    )
    _add_table(
        doc,
        ["Semana", "Actividades principales", "Horas acum."],
        [
            ["1", "Infraestructura, login, inicio consulta afiliado", "40"],
            ["2", "Consulta afiliado, análisis Messiah paso 1", "72"],
            ["3", "Direccionamiento y validación medicamentos", "112"],
            ["4", "Repositorio solicitud, endpoint paso 1", "152"],
            ["5", "Soporte PDF paso 1, pruebas paso 1", "184"],
            ["6", "Análisis paso 2, emisión autorización", "216"],
            ["7", "Activación, confirmación, preferencias", "248"],
            ["8", "Contabilización NC-AT, concurrencia", "280"],
            ["9", "PDF Jasper, caché, idempotencia", "312"],
            ["10", "Pruebas integrales, manual, despliegue", "360"],
        ],
    )

    doc.add_heading("9. Supuestos y exclusiones", level=1)
    doc.add_paragraph(
        "Incluido en la estimación:\n"
        "• Desarrollo backend Python/FastAPI con integración directa a BD Messiah.\n"
        "• Alineación funcional con flujos Messiah existentes (no modificación del código Java Messiah).\n"
        "• Documentación de consumo y colección Postman.\n\n"
        "Excluido o fuera de alcance (horas adicionales):\n"
        "• Desarrollo frontend o cambios en la aplicación web Messiah.\n"
        "• Migraciones de esquema BD (solo INSERT/UPDATE sobre tablas existentes).\n"
        "• Certificación de seguridad / pentest.\n"
        "• Soporte post-producción prolongado (se cotiza aparte).\n"
        "• Configuración SFTP Messiah en infraestructura del cliente (solo integración en código).\n\n"
        "Riesgos que pueden extender el cronograma (+15–25 %):\n"
        "• Cambios de reglas de negocio no documentadas en Messiah.\n"
        "• Incompatibilidad de versión Jasper/Java en el servidor destino.\n"
        "• Datos maestros incompletos (tarifarios, concepto nota técnica, sc_contabiliza)."
    )

    doc.add_heading("10. Archivos de referencia en el repositorio", level=1)
    _add_table(
        doc,
        ["Área", "Archivo principal"],
        [
            ["Rutas API", "app/api/routes.py"],
            ["Login / JWT", "app/core/security.py"],
            ["Consulta afiliado", "app/repositories/postgres_repository.py"],
            ["Paso 1 y 2", "app/services/autorizacion_medicamentos_ips_service.py"],
            ["Direccionamiento", "app/services/direccionamiento_service.py"],
            ["Persistencia Messiah", "app/repositories/messiah_direccionamiento_repository.py"],
            ["Contabilización", "app/repositories/messiah_contabilizacion_repository.py"],
            ["PDF Jasper", "app/services/messiah_autorizacion_pdf_service.py"],
            ["Manual consumo", "Manual_Consumo_Ordenamiento_Medicamentos.docx"],
            ["Postman", "postman/EPS Familiar de Colombia API - GHG.postman_collection.json"],
        ],
    )

    doc.save(out)
    print(f"Escrito: {out}")


if __name__ == "__main__":
    main()
